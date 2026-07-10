"""DOCX ingestion: body-order extraction of text and table chunks.

Phase 1 scope: text paragraphs and tables (no embedded images). Output uses the
same ``ExtractedChunk`` contract as the PDF path so embedding, indexing,
retrieval, and computed charts stay format-agnostic.
"""

from __future__ import annotations

import logging

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ingestion.chunking import chunk_text, normalize_whitespace
from app.ingestion.models import ExtractedChunk
from app.ingestion.tables import table_signature, table_to_markdown

logger = logging.getLogger(__name__)

MIN_TEXT_WORDS = 8
MIN_TABLE_WORDS = 4


def _iter_block_items(document):
    """Yield paragraphs and tables in document reading order.

    ``document.paragraphs`` and ``document.tables`` are separate lists that lose
    interleaving, so we walk the body XML directly. Nested tables live inside
    cells (not the body), so they are not double-counted here.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or "" if paragraph.style else ""
    except Exception:
        return ""


def _is_heading(style_name: str) -> bool:
    normalized = style_name.strip().lower()
    return normalized == "title" or normalized.startswith("heading")


def _table_rows(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([normalize_whitespace(cell.text) for cell in row.cells])
    return rows


def _base_metadata(block_index: int, section: str | None) -> dict:
    extra: dict = {"source_format": "docx", "block_index": block_index}
    if section:
        extra["section"] = section
    return extra


def _text_chunks(
    paragraph: Paragraph,
    block_index: int,
    section: str | None,
) -> list[ExtractedChunk]:
    text = normalize_whitespace(paragraph.text)
    if not text:
        return []

    chunks: list[ExtractedChunk] = []
    for piece in chunk_text(text):
        if len(piece.split()) < MIN_TEXT_WORDS:
            continue
        chunks.append(
            ExtractedChunk(
                content=piece,
                page_number=block_index,
                chunk_type="text",
                extraction_method="docx_native",
                extra_metadata=_base_metadata(block_index, section),
            )
        )
    return chunks


def _table_chunk(
    rows: list[list[str]],
    block_index: int,
    section: str | None,
) -> ExtractedChunk | None:
    markdown = table_to_markdown(rows)
    if not markdown or len(markdown.split()) < MIN_TABLE_WORDS:
        return None

    extra = _base_metadata(block_index, section)
    _, headers = table_signature(rows)
    if headers:
        extra["table_headers"] = list(headers)

    return ExtractedChunk(
        content=markdown,
        page_number=block_index,
        chunk_type="table",
        extraction_method="docx_native",
        extra_metadata=extra,
    )


def extract_docx_chunks(
    docx_path: str,
    *,
    doc_id: str | None = None,
    user_id: int | None = None,
) -> list[ExtractedChunk]:
    """Extract text and table chunks from a DOCX file in reading order.

    ``page_number`` carries a 1-based block ordinal (position of the source
    paragraph/table in the document) since DOCX has no fixed pagination. The
    most recent heading/title is tracked and stored as ``extra_metadata.section``
    to give retrieval and citations lightweight structural context.
    """
    document = Document(docx_path)
    chunks: list[ExtractedChunk] = []
    block_index = 0
    current_section: str | None = None

    for block in _iter_block_items(document):
        block_index += 1
        if isinstance(block, Paragraph):
            if _is_heading(_style_name(block)):
                heading_text = normalize_whitespace(block.text)
                if heading_text:
                    current_section = heading_text
            chunks.extend(_text_chunks(block, block_index, current_section))
        elif isinstance(block, Table):
            table_chunk = _table_chunk(_table_rows(block), block_index, current_section)
            if table_chunk is not None:
                chunks.append(table_chunk)

    logger.info("Extracted %s chunks from DOCX doc_id=%s", len(chunks), doc_id)
    return chunks
