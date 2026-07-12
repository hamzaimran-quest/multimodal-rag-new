"""Unit tests for DOCX embedded image extraction."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

from app.config import settings
from app.ingestion.docx_images import extract_docx_image_chunks


def _write_docx_with_image(path: Path) -> None:
    doc = Document()
    doc.add_paragraph(
        "Introduction paragraph with enough words to provide nearby context for image retrieval."
    )
    paragraph = doc.add_paragraph()
    img = Image.new("RGB", (120, 120), color=(200, 40, 40))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    paragraph.add_run().add_picture(buf, width=Inches(1.2))
    doc.add_paragraph(
        "Caption paragraph describing the red square diagram used for testing purposes here."
    )
    doc.save(path)


def test_extract_docx_image_chunks_persists_and_indexes(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    _write_docx_with_image(docx_path)

    chunks = extract_docx_image_chunks(
        str(docx_path), doc_id="doc-1", user_id=7
    )

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.chunk_type == "image"
    assert chunk.extraction_method == "docx_ocr_proxy"
    assert chunk.page_number == 2
    assert chunk.extra_metadata["source_format"] == "docx"
    assert chunk.extra_metadata["block_index"] == 2
    assert chunk.bbox is None
    assert chunk.image_path is not None
    assert "Nearby text" in chunk.content or "red square" in chunk.content.lower()

    image_file = settings.resolved_images_dir / "7" / "doc-1" / "block2_img0.png"
    assert image_file.is_file()


def test_extract_docx_image_chunks_skips_without_ids(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    _write_docx_with_image(docx_path)
    assert extract_docx_image_chunks(str(docx_path)) == []
